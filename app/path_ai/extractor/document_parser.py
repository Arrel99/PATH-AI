from pathlib import Path
from app.path_ai.monitoring.logger import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def parse_pdf(file_path: str) -> dict:
    import fitz  # PyMuPDF
    from app.path_ai.extractor.ocr_processor import ocr_image

    doc = fitz.open(file_path)
    text_parts = []
    images = []
    page_count = len(doc)

    for page_num, page in enumerate(doc):
        # Extract text
        page_text = page.get_text("text")
        if page_text.strip():
            text_parts.append(page_text.strip())

        # Extract images
        for img_index, img in enumerate(page.get_images(full=True)):
            try:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                images.append(image_bytes)

                # OCR the image
                ocr_text = ocr_image(image_bytes)
                if ocr_text:
                    text_parts.append(f"[Gambar halaman {page_num + 1}]\n{ocr_text}")
            except Exception as e:
                logger.warning(
                    "Failed to extract image",
                    page=page_num + 1,
                    error=str(e),
                )

    doc.close()

    combined_text = "\n\n".join(text_parts)
    logger.info(
        "PDF parsed",
        path=file_path,
        pages=page_count,
        text_length=len(combined_text),
        images_found=len(images),
    )

    return {
        "text": combined_text,
        "images": images,
        "page_count": page_count,
        "source": Path(file_path).name,
    }


def parse_docx(file_path: str) -> dict:
    from docx import Document
    from app.path_ai.extractor.ocr_processor import ocr_image

    doc = Document(file_path)
    text_parts = []
    images = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    # Extract images
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_bytes = rel.target_part.blob
                images.append(image_bytes)

                ocr_text = ocr_image(image_bytes)
                if ocr_text:
                    text_parts.append(f"[Gambar]\n{ocr_text}")
            except Exception as e:
                logger.warning("Failed to extract docx image", error=str(e))

    combined_text = "\n\n".join(text_parts)
    logger.info(
        "DOCX parsed",
        path=file_path,
        text_length=len(combined_text),
        images_found=len(images),
    )

    return {
        "text": combined_text,
        "images": images,
        "page_count": len(doc.paragraphs),
        "source": Path(file_path).name,
    }


def parse_file(file_path: str) -> dict:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_path)
